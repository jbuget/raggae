import sys
import types
from io import BytesIO

import pytest
from docx import Document as DocxDocument

from raggae.domain.exceptions.document_exceptions import DocumentExtractionError
from raggae.infrastructure.services.multiformat_document_text_extractor import (
    MultiFormatDocumentTextExtractor,
)


class TestMultiFormatDocumentTextExtractor:
    @pytest.fixture
    def extractor(self) -> MultiFormatDocumentTextExtractor:
        return MultiFormatDocumentTextExtractor()

    async def test_extract_text_txt_success(
        self,
        extractor: MultiFormatDocumentTextExtractor,
    ) -> None:
        # Given
        content = b"hello\nworld"

        # When
        result = await extractor.extract_text(
            file_name="notes.txt",
            content=content,
            content_type="text/plain",
        )

        # Then
        assert result == "hello\nworld"

    async def test_extract_text_md_normalizes_whitespace(
        self,
        extractor: MultiFormatDocumentTextExtractor,
    ) -> None:
        # Given
        content = b"  # Title\r\n\r\ncontent   "

        # When
        result = await extractor.extract_text(
            file_name="README.md",
            content=content,
            content_type="text/markdown",
        )

        # Then
        assert result == "# Title\n\ncontent"

    async def test_extract_text_pdf_uses_pdf_extractor(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        monkeypatch.setattr(extractor, "_extract_pdf", lambda _: "pdf text")

        # When
        result = await extractor.extract_text(
            file_name="file.pdf",
            content=b"%PDF-1.7",
            content_type="application/pdf",
        )

        # Then
        assert result == "pdf text"

    async def test_extract_pdf_includes_page_markers(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given — fake pdfplumber with two pages, no tables
        class _FakePage:
            def __init__(self, text: str) -> None:
                self._text = text

            def find_tables(self) -> list[object]:
                return []

            def extract_text(self) -> str:
                return self._text

            def filter(self, fn: object) -> "_FakePage":
                return self

        class _FakePdf:
            pages = [_FakePage("page one"), _FakePage("page two")]

            def __enter__(self) -> "_FakePdf":
                return self

            def __exit__(self, *args: object) -> None:
                pass

        fake_pdfplumber = types.SimpleNamespace(open=lambda buf: _FakePdf())
        monkeypatch.setitem(sys.modules, "pdfplumber", fake_pdfplumber)
        monkeypatch.setattr(extractor._pdf_table_extractor, "extract_tables", lambda _: [])

        # When
        result = extractor._extract_pdf(b"%PDF-1.7")

        # Then
        assert "[[PAGE:1]]" in result
        assert "page one" in result
        assert "[[PAGE:2]]" in result
        assert "page two" in result

    async def test_extract_pdf_table_chunk_appended_after_text(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given — PdfTableExtractor returns one table on page 1
        from raggae.infrastructure.services.pdf_table_extractor import TableData

        fake_table = TableData(
            rows=[["Produit", "Prix"], ["Widget A", "48 €"]],
            page_number=1,
            table_index=1,
            bbox=(10.0, 20.0, 200.0, 100.0),
        )

        class _FakePage:
            def find_tables(self) -> list[object]:
                return []

            def extract_text(self) -> str:
                return "some text"

            def filter(self, fn: object) -> "_FakePage":
                return self

        class _FakePdf:
            pages = [_FakePage()]

            def __enter__(self) -> "_FakePdf":
                return self

            def __exit__(self, *args: object) -> None:
                pass

        monkeypatch.setitem(sys.modules, "pdfplumber", types.SimpleNamespace(open=lambda buf: _FakePdf()))
        monkeypatch.setattr(extractor._pdf_table_extractor, "extract_tables", lambda _: [fake_table])

        # When
        result = extractor._extract_pdf(b"%PDF-1.7")

        # Then
        assert "[[PAGE:1]] [TABLE:1]" in result
        assert "| Produit | Prix |" in result
        assert "| Widget A | 48 € |" in result

    async def test_extract_pdf_table_bbox_filter_applied_to_text(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given — filter() is called when table bbox is present, returning filtered text
        from raggae.infrastructure.services.pdf_table_extractor import TableData

        fake_table = TableData(
            rows=[["H"], ["v"]],
            page_number=1,
            table_index=1,
            bbox=(10.0, 20.0, 200.0, 100.0),
        )

        class _FakeFilteredPage:
            def extract_text(self) -> str:
                return "only prose here"

        class _FakePage:
            def find_tables(self) -> list[object]:
                return []

            def extract_text(self) -> str:
                return "prose + TABLE CONTENT"

            def filter(self, fn: object) -> _FakeFilteredPage:
                return _FakeFilteredPage()

        class _FakePdf:
            pages = [_FakePage()]

            def __enter__(self) -> "_FakePdf":
                return self

            def __exit__(self, *args: object) -> None:
                pass

        monkeypatch.setitem(sys.modules, "pdfplumber", types.SimpleNamespace(open=lambda buf: _FakePdf()))
        monkeypatch.setattr(extractor._pdf_table_extractor, "extract_tables", lambda _: [fake_table])

        # When
        result = extractor._extract_pdf(b"%PDF-1.7")

        # Then — filtered text used, raw text not present
        assert "only prose here" in result
        assert "TABLE CONTENT" not in result

    async def test_extract_pdf_table_extraction_failure_still_returns_text(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given — PdfTableExtractor raises; text extraction must still succeed
        class _FakePage:
            def find_tables(self) -> list[object]:
                return []

            def extract_text(self) -> str:
                return "fallback text"

            def filter(self, fn: object) -> "_FakePage":
                return self

        class _FakePdf:
            pages = [_FakePage()]

            def __enter__(self) -> "_FakePdf":
                return self

            def __exit__(self, *args: object) -> None:
                pass

        monkeypatch.setitem(sys.modules, "pdfplumber", types.SimpleNamespace(open=lambda buf: _FakePdf()))
        monkeypatch.setattr(
            extractor._pdf_table_extractor,
            "extract_tables",
            lambda _: (_ for _ in ()).throw(RuntimeError("pdfplumber crashed")),
        )

        # When
        result = extractor._extract_pdf(b"%PDF-1.7")

        # Then — text still extracted despite table extractor failure
        assert "[[PAGE:1]]" in result
        assert "fallback text" in result

    async def test_extract_text_docx_uses_docx_extractor(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        monkeypatch.setattr(extractor, "_extract_docx", lambda _: "docx text")

        # When
        result = await extractor.extract_text(
            file_name="file.docx",
            content=b"PK...",
            content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        )

        # Then
        assert result == "docx text"

    async def test_extract_text_docx_extracts_paragraphs_and_tables(
        self,
        extractor: MultiFormatDocumentTextExtractor,
    ) -> None:
        # Given
        document = DocxDocument()
        document.add_paragraph("Transcript intro")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Action item"
        table.cell(0, 1).text = "Prepare demo for next meeting"
        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()

        # When
        result = await extractor.extract_text(
            file_name="meeting.docx",
            content=content,
            content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        )

        # Then
        assert "Transcript intro" in result
        assert "Action item" in result
        assert "Prepare demo for next meeting" in result

    async def test_extract_text_doc_raises_not_supported_error(
        self,
        extractor: MultiFormatDocumentTextExtractor,
    ) -> None:
        # When / Then
        with pytest.raises(DocumentExtractionError):
            await extractor.extract_text(
                file_name="legacy.doc",
                content=b"DOC",
                content_type="application/msword",
            )

    async def test_extract_text_unknown_extension_raises_error(
        self,
        extractor: MultiFormatDocumentTextExtractor,
    ) -> None:
        # When / Then
        with pytest.raises(DocumentExtractionError):
            await extractor.extract_text(
                file_name="file.bin",
                content=b"binary",
                content_type="application/octet-stream",
            )

    async def test_extract_text_pptx_uses_pptx_extractor(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        monkeypatch.setattr(extractor, "_extract_pptx", lambda _: "[SLIDE:1]\n# Title\n\nContent.")

        # When
        result = await extractor.extract_text(
            file_name="deck.pptx",
            content=b"PK...",
            content_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        )

        # Then
        assert "[SLIDE:1]" in result
        assert "Title" in result

    async def test_extract_pptx_includes_slide_markers(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        import sys
        import types
        from io import BytesIO

        class _FakeTextFrame:
            def __init__(self, text: str) -> None:
                self.text = text

        class _FakeNoteSlide:
            notes_text_frame = _FakeTextFrame("")

        class _FakeShape:
            def __init__(self, text: str, top: int = 0, left: int = 0) -> None:
                self.text_frame = _FakeTextFrame(text)
                self.has_text_frame = True
                self.has_table = False
                self.top = top
                self.left = left

        class _FakeShapes:
            def __init__(self, shapes: list[_FakeShape], title: _FakeShape | None) -> None:
                self._shapes = shapes
                self.title = title

            def __iter__(self):  # type: ignore[override]
                return iter(self._shapes)

        class _FakeSlide:
            def __init__(self, title_text: str, body_text: str) -> None:
                title_shape = _FakeShape(title_text)
                body_shape = _FakeShape(body_text, top=100)
                self.shapes = _FakeShapes([title_shape, body_shape], title=title_shape)
                self.notes_slide = _FakeNoteSlide()

        class _FakePresentation:
            def __init__(self, _buffer: BytesIO) -> None:
                self.slides = [
                    _FakeSlide("Slide One", "Body of slide one."),
                    _FakeSlide("Slide Two", "Body of slide two."),
                ]
                self.core_properties = types.SimpleNamespace(title=None, author=None, created=None)

        fake_pptx = types.SimpleNamespace(Presentation=_FakePresentation)
        monkeypatch.setitem(sys.modules, "pptx", fake_pptx)

        # When
        result = extractor._extract_pptx(b"PK...")

        # Then
        assert "[SLIDE:1]" in result
        assert "[SLIDE:2]" in result
        assert "Slide One" in result
        assert "Slide Two" in result

    async def test_extract_pptx_includes_speaker_notes(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        import sys
        import types
        from io import BytesIO

        class _FakeTextFrame:
            def __init__(self, text: str) -> None:
                self.text = text

        class _FakeNoteSlide:
            def __init__(self, notes: str) -> None:
                self.notes_text_frame = _FakeTextFrame(notes)

        class _FakeShape:
            def __init__(self, text: str) -> None:
                self.text_frame = _FakeTextFrame(text)
                self.has_text_frame = True
                self.has_table = False
                self.top = 0
                self.left = 0

        class _FakeShapes:
            def __init__(self, shape: _FakeShape) -> None:
                self._shapes = [shape]
                self.title = shape

            def __iter__(self):  # type: ignore[override]
                return iter(self._shapes)

        class _FakeSlide:
            def __init__(self) -> None:
                shape = _FakeShape("Title")
                self.shapes = _FakeShapes(shape)
                self.notes_slide = _FakeNoteSlide("These are my speaker notes.")

        class _FakePresentation:
            def __init__(self, _buffer: BytesIO) -> None:
                self.slides = [_FakeSlide()]

        fake_pptx = types.SimpleNamespace(Presentation=_FakePresentation)
        monkeypatch.setitem(sys.modules, "pptx", fake_pptx)

        # When
        result = extractor._extract_pptx(b"PK...")

        # Then
        assert "[NOTES]" in result
        assert "These are my speaker notes." in result

    async def test_extract_pptx_sorts_shapes_by_position(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        import sys
        import types
        from io import BytesIO

        class _FakeTextFrame:
            def __init__(self, text: str) -> None:
                self.text = text

        class _FakeNoteSlide:
            notes_text_frame = _FakeTextFrame("")

        class _FakeShape:
            def __init__(self, text: str, top: int, left: int) -> None:
                self.text_frame = _FakeTextFrame(text)
                self.has_text_frame = True
                self.has_table = False
                self.top = top
                self.left = left

        class _FakeShapes:
            def __init__(self) -> None:
                # Intentionally unordered
                self._shapes = [
                    _FakeShape("Third", top=200, left=0),
                    _FakeShape("First", top=0, left=0),
                    _FakeShape("Second", top=100, left=0),
                ]
                self.title = None

            def __iter__(self):  # type: ignore[override]
                return iter(self._shapes)

        class _FakeSlide:
            def __init__(self) -> None:
                self.shapes = _FakeShapes()
                self.notes_slide = _FakeNoteSlide()

        class _FakePresentation:
            def __init__(self, _buffer: BytesIO) -> None:
                self.slides = [_FakeSlide()]

        fake_pptx = types.SimpleNamespace(Presentation=_FakePresentation)
        monkeypatch.setitem(sys.modules, "pptx", fake_pptx)

        # When
        result = extractor._extract_pptx(b"PK...")

        # Then - "First" must appear before "Third"
        assert result.index("First") < result.index("Third")

    async def test_extract_pptx_skips_empty_shapes(
        self,
        extractor: MultiFormatDocumentTextExtractor,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Given
        import sys
        import types
        from io import BytesIO

        class _FakeTextFrame:
            def __init__(self, text: str) -> None:
                self.text = text

        class _FakeNoteSlide:
            notes_text_frame = _FakeTextFrame("")

        class _FakeShape:
            def __init__(self, text: str) -> None:
                self.text_frame = _FakeTextFrame(text)
                self.has_text_frame = True
                self.has_table = False
                self.top = 0
                self.left = 0

        class _FakeShapes:
            def __init__(self) -> None:
                self._shapes = [_FakeShape(""), _FakeShape("Real content.")]
                self.title = None

            def __iter__(self):  # type: ignore[override]
                return iter(self._shapes)

        class _FakeSlide:
            def __init__(self) -> None:
                self.shapes = _FakeShapes()
                self.notes_slide = _FakeNoteSlide()

        class _FakePresentation:
            def __init__(self, _buffer: BytesIO) -> None:
                self.slides = [_FakeSlide()]

        fake_pptx = types.SimpleNamespace(Presentation=_FakePresentation)
        monkeypatch.setitem(sys.modules, "pptx", fake_pptx)

        # When
        result = extractor._extract_pptx(b"PK...")

        # Then
        assert "Real content." in result
        # Empty shape text not present as isolated blank line block
        lines = [ln for ln in result.split("\n") if ln.strip()]
        assert all(ln.strip() for ln in lines)

    async def test_extract_text_ppt_raises_error(
        self,
        extractor: MultiFormatDocumentTextExtractor,
    ) -> None:
        # When / Then
        with pytest.raises(DocumentExtractionError):
            await extractor.extract_text(
                file_name="legacy.ppt",
                content=b"binary",
                content_type="application/vnd.ms-powerpoint",
            )
